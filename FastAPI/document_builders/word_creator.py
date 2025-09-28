from docx import Document
from docx.shared import Inches
import json
import re

def save_to_word(content, filename, chart_metadata=None):
    """
    Export report content to a Word file, embedding charts inline
    at their placeholders like [[FIGURE:chart1]].
    """
    if isinstance(content, str):
        content = content.strip()
        if not content:
            raise ValueError("Draft content is empty")

        # Remove markdown code fences
        content = re.sub(r"^```(?:json)?|```$", "", content.strip(), flags=re.MULTILINE).strip()
        content = json.loads(content)

    # Map chart IDs to paths
    chart_map = {c["id"]: c["path"] for c in (chart_metadata or [])}

    doc = Document()
    doc.add_heading(content.get("title", "Automotive Supply Chain Report"), 0)

    def insert_text_with_figures(text):
        # Split by figure placeholders
        parts = re.split(r"(\[\[FIGURE:([^\]]+)\]\])", text)
        for i, part in enumerate(parts):
            if part.startswith("[[FIGURE:"):
                # It's a placeholder
                fig_id = re.match(r"\[\[FIGURE:([^\]]+)\]\]", part).group(1)
                if fig_id in chart_map:
                    try:
                        doc.add_picture(chart_map[fig_id], width=Inches(5))
                        doc.add_paragraph(f"Figure: {fig_id}")
                    except Exception:
                        doc.add_paragraph(f"(Missing figure: {fig_id})")
                else:
                    doc.add_paragraph(f"(No chart found for {fig_id})")
            elif part.strip():
                doc.add_paragraph(part)

    def insert_bullets(doc, bullets):
        for b in bullets:
            para = doc.add_paragraph(b)
            para.style = 'ListBullet'

    for section in content.get("sections", []):
        doc.add_heading(section.get("heading", "No Heading"), level=1)
        insert_text_with_figures(section.get("content", ""))

        # Process figures field if it exists
        figures = section.get("figures", [])
        if figures:
            for figure in figures:
                insert_text_with_figures(figure)

        insert_bullets(doc, section.get("bullet_points", []))
        for subsection in section.get("subsections", []):
            doc.add_heading(subsection.get("heading", "No Subheading"), level=2)
            insert_text_with_figures(subsection.get("content", ""))

            # Process figures field for subsections if it exists
            subsection_figures = subsection.get("figures", [])
            if subsection_figures:
                for figure in subsection_figures:
                    insert_text_with_figures(figure)

    doc.save(filename)
    return filename